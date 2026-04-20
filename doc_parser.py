from pdf2image import convert_from_path
import os
import re
import datetime
from custom_exceptions import ExtractionTimeOut
from docx import Document

OCR_PROMPT = """
You are a document-to-text conversion engine for building a searchable knowledge base.

Convert the provided document into clean, structured, chunk-ready plain text.

Requirements:

Process the document page by page.

For each page, output in the exact format:
{extracted_text}

STRICT RULES:
- Output ONLY the extracted/converted content text.
- Do NOT describe the image.
- Do NOT summarise.
- Do NOT explain.
- Do NOT add commentary.
- Do NOT guess or hallucinate missing information.
- Preserve wording and numbers exactly as written.
- If text is unreadable, omit it.
- Preserve capitalization of headings.

SYMBOL CONVERSIONS:
- Checkmark/tick → "✓"
- Cross → "✗"
- Arrow → "→"
- Currency, percentages, dates → preserve exactly

GENERAL STRUCTURE RULES:
1) Normal Paragraph Pages:
- Preserve paragraphs.
- Preserve bullet points using "-" prefix.
- Keep headings separated by blank lines.

2) Comparison Layouts (tables, pricing grids, feature matrices):
- Do NOT output as rows/columns.
- Rewrite into separate self-contained sections per item/plan/package.
- Each section must include all associated facts (prices, quantities, features, limits).
- Format:
  [ITEM NAME]
  - fact
  - fact
  - includes: ...
  - price: ...

3) Process / Flow / Funnel Diagrams:
- Detect directional or sequential relationships.
- Convert to explicit ordered sequence using "→".
- Include descriptions under each step.
- Format:
  PROCESS:
  Step 1: ...
  - description
  →
  Step 2: ...
  - description
  →
  Step 3: ...

4) Step-by-Step Instructions / Roadmaps:
- Preserve order.
- Use numbered format:
  STEP 1: ...
  - details
  STEP 2: ...
  - details

5) Metrics / Results / Case Studies:
- Keep numeric results clearly tied to their labels.
- Format:
  [METRIC TITLE]
  - value: ...
  - timeframe: ...
  - context: ...

6) Multi-Column Marketing Layouts:
- Ignore visual column structure.
- Merge into logical reading order.
- Preserve section headings.

7) FAQs:
- Format clearly as:
  Q: ...
  A: ...

8) Testimonials / Quotes:
- Preserve quoted text.
- Attribute if visible.

9) Contact Information:
- Keep structured as:
  CONTACT:
  - phone:
  - email:
  - website:

FINAL REQUIREMENTS:
- Each logical section must be self-contained.
- Separate sections with a blank line.
- Ensure output is readable as standalone text without needing the original layout.
"""

STRUCTURED_OCR_PROMPT = OCR_PROMPT.replace(
    "Requirements:",
    "HEADING DETECTION RULES:\n"
    "When you encounter a section-level heading, output it on its own line with [HEADING] "
    "immediately before the heading text, like this:\n"
    "[HEADING] Section Title\n\n"
    "There is usually 1-3 section headings in a page \n"
    "A heading IS:\n"
    "- A short title or label (1–8 words) that introduces a new topic or section\n"
    "- Visually distinct in the original document (bold, larger font, underlined, or on its own line)\n"
    "- NOT a full sentence and does NOT end with punctuation (. , : ; ? !)\n\n"
    "A heading is NOT:\n"
    "- A bullet point label or sub-item\n"
    "- Regular body text or a sentence within a paragraph\n"
    "- A caption, table header, or form label\n\n"
    "EXAMPLES:\n\n"
    "Example 1 — Service brochure:\n"
    "[HEADING] Our Services\n"
    "We provide a full range of digital marketing services tailored to your business goals.\n\n"
    "[HEADING] Search Engine Optimisation\n"
    "Our SEO team improves your website's visibility on Google through technical audits, "
    "keyword research, and content strategy.\n\n"
    "[HEADING] Paid Advertising\n"
    "We manage Google Ads and Meta campaigns to drive targeted traffic and maximise ROI.\n\n"
    "Example 2 — Pricing page:\n"
    "[HEADING] Pricing Plans\n"
    "Choose the package that best fits your needs. All plans include a dedicated account manager.\n\n"
    "[HEADING] Starter Package\n"
    "- Monthly management fee: £500\n"
    "- Includes: SEO audit, 2 blog posts, monthly report\n\n"
    "[HEADING] Growth Package\n"
    "- Monthly management fee: £1,200\n"
    "- Includes: Full SEO, Google Ads management, weekly reporting\n\n"
    "Example 3 — Company overview:\n"
    "[HEADING] About Us\n"
    "Suri Marketing was founded in 2018 in Birmingham with a mission to help local businesses grow online.\n\n"
    "[HEADING] Our Approach\n"
    "We take a data-driven approach to every campaign, combining creative content with measurable results.\n\n"
    "[HEADING] Why Choose Us\n"
    "- Over 150 clients across the UK\n"
    "- Average 3x ROI on paid campaigns\n"
    "- Dedicated account manager for every client\n\n"
    "Example 4 — FAQ page:\n"
    "[HEADING] Frequently Asked Questions\n\n"
    "Q: How long does it take to see SEO results?\n"
    "A: Most clients see measurable improvements within 3–6 months.\n\n"
    "Q: Do you offer one-off projects?\n"
    "A: Yes, we offer one-off audits and strategy sessions alongside ongoing retainers.\n\n"
    "Requirements:",
    1,
)

# Short connector words used in plain-text heading heuristic
_CONNECTORS = {"a", "an", "the", "of", "in", "on", "at", "to", "for", "and", "or", "but", "with", "by"}


def _is_plain_text_heading(line):
    """Return True if line looks like a plain-text section heading."""
    line = line.strip()
    if not line or line[-1] in ".,:;?!":
        return False
    words = line.split()
    if len(words) > 10:
        return False
    for word in words:
        if word.lower() in _CONNECTORS:
            continue
        if not word[0].isupper():
            return False
    return True


def _is_all_caps_heading(line):
    """Return True if line is ALL CAPS words (likely a plain-text heading)."""
    line = line.strip()
    if not line:
        return False
    words = line.split()
    if len(words) > 10:
        return False
    # Allow digits and common symbols (e.g. "SECTION 1", "SERVICES & PRICING")
    return all(w.isupper() for w in words if w.isalpha())


def __create_file(client, file_path):
    with open(file_path, "rb") as f:
        result = client.files.create(file=f, purpose="vision")
    return result.id

def extract_from_pdf(client, pdf_file_path, images_dir_name, ocr_prompt=None):
    start_time = datetime.datetime.now()
    max_time = 300
    prompt = ocr_prompt if ocr_prompt is not None else OCR_PROMPT

    images = convert_from_path(pdf_file_path)

    doc_name = os.path.splitext(os.path.basename(pdf_file_path))[0]
    directory_name = os.path.join(images_dir_name, doc_name)
    os.makedirs(directory_name, exist_ok=True)

    all_extracted_texts = []

    for page_no, img in enumerate(images, start=1):
        delta_time = datetime.datetime.now() - start_time

        if delta_time.total_seconds() > max_time:
            raise ExtractionTimeOut("Document extraction exceeded max time allowed.")

        file_path = os.path.join(directory_name, f"page{page_no}.png")
        img.save(file_path, "PNG")

        file_id = None
        try:
            file_id = __create_file(client, file_path)

            response = client.responses.create(
                model="gpt-4.1-mini",
                input=[
                    {
                        "role": "system",
                        "content": "You are a document-to-text conversion engine.",
                    },
                    {
                        "role": "user",
                        "content": [
                            {"type": "input_text", "text": prompt},
                            {"type": "input_image", "file_id": file_id},
                        ],
                    },
                ],
            )

            # If your SDK supports it, this is fine:
            text = getattr(response, "output_text", None)
            if text is None:
                text = str(response)
            all_extracted_texts.append(text)

        except Exception as e:
            all_extracted_texts.append(
                f"--- PAGE {page_no} ---\n[ERROR extracting page: {e}]"
            )
        finally:
            if file_id is not None:
                try:
                    client.files.delete(file_id)
                except Exception:
                    pass

    combined = "\n\n".join(all_extracted_texts)

    # If structured mode was requested but LLM didn't add any [HEADING] markers,
    # fall back to plain-text heuristics on the extracted content.
    if ocr_prompt is not None and "[HEADING]" not in combined:
        print("[doc_parser] No [HEADING] markers in PDF LLM output — applying plain-text heuristics.")
        combined = _apply_plain_text_heading_heuristics(combined)

    return combined

def extract_from_image(client, file_path, ocr_prompt=None):
    prompt = ocr_prompt if ocr_prompt is not None else OCR_PROMPT
    file_id = None
    try:
        file_id = __create_file(client, file_path)

        response = client.responses.create(
            model="gpt-4.1-mini",
            input=[
                {
                    "role": "system",
                    "content": "You are a document-to-text conversion engine.",
                },
                {
                    "role": "user",
                    "content": [
                        {"type": "input_text", "text": prompt},
                        {"type": "input_image", "file_id": file_id},
                    ],
                },
            ],
        )

        text = getattr(response, "output_text", None)
        if text is None:
            text = str(response)
        return text.strip()

    finally:
        if file_id is not None:
            try:
                client.files.delete(file_id)
            except Exception:
                pass


def extract_from_docx(file_path):
    doc = Document(file_path)
    lines = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    return "\n\n".join(lines)


def extract_from_docx_structured(file_path):
    doc = Document(file_path)
    lines = []
    heading_found = False

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style.name.startswith("Heading"):
            lines.append(f"[HEADING] {text}")
            heading_found = True
        else:
            lines.append(text)

    if not heading_found:
        print(f"[doc_parser] WARNING: No heading-style paragraphs found in {file_path}. "
              "Falling back to plain-text heuristics.")
        return _apply_plain_text_heading_heuristics("\n\n".join(lines))

    return "\n\n".join(lines)


def _apply_plain_text_heading_heuristics(text):
    """Apply ALL CAPS and title-case heuristics to mark headings in plain text."""
    lines = text.splitlines()
    result = []
    heading_found = False

    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            result.append(line)
            continue

        next_line = lines[i + 1].strip() if i + 1 < len(lines) else ""
        is_heading = (
            stripped.startswith("#")
            or _is_all_caps_heading(stripped)
            or (_is_plain_text_heading(stripped) and next_line == "")
        )

        if is_heading:
            clean = stripped.lstrip("#").strip()
            result.append(f"[HEADING] {clean}")
            heading_found = True
        else:
            result.append(line)

    if not heading_found:
        print("[doc_parser] WARNING: No headings detected by any heuristic. "
              "Document will be treated as a single parent section.")

    return "\n".join(result)


def extract_from_text_file(file_path):
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


def extract_from_md(file_path):
    return extract_from_text_file(file_path)


def extract_from_text_file_structured(file_path):
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()
    return _apply_plain_text_heading_heuristics(text)


def extract_from_md_structured(file_path):
    return extract_from_text_file_structured(file_path)


def extract_doc_info(client, file_path, structured=False):
    ext = os.path.splitext(file_path)[1].lower()
    print("file type: ", ext)

    if ext == ".pdf":
        images_dir = os.path.join(os.path.dirname(file_path), "extracted_images")
        prompt = STRUCTURED_OCR_PROMPT if structured else None
        return extract_from_pdf(client, file_path, images_dir, ocr_prompt=prompt)
    elif ext in (".png", ".jpg", ".jpeg"):
        prompt = STRUCTURED_OCR_PROMPT if structured else None
        return extract_from_image(client, file_path, ocr_prompt=prompt)
    elif ext == ".docx":
        return extract_from_docx_structured(file_path) if structured else extract_from_docx(file_path)
    elif ext == ".txt":
        return extract_from_text_file_structured(file_path) if structured else extract_from_text_file(file_path)
    elif ext == ".md":
        return extract_from_md_structured(file_path) if structured else extract_from_md(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")